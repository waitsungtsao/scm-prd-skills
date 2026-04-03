#!/usr/bin/env python3
"""
md2docx.py — PRD Markdown → Word 文档转换（含图片嵌入）

维护状态: 降级备选方案。推荐使用 md2docx.mjs（Node.js 版，排版精度更高）。
本文件仅在用户无 Node.js 环境时作为 fallback 使用，不再追求与 JS 版功能对齐。

将 PRD Markdown 文件转换为 Word (.docx) 格式，自动嵌入 diagrams/ 目录下的 PNG 图片。

用法:
    python md2docx.py <PRD.md文件路径>
    python md2docx.py requirements/REQ-xxx/PRD-xxx.md

依赖: python-docx (pip install python-docx)
兼容: Python 3.8+
"""

import sys
import os
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print(
        "错误: 需要 python-docx 库。请运行: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# =============================================================================
# 常量
# =============================================================================

MONOSPACE_FONT = "Consolas"
MONOSPACE_SIZE = Pt(10)
CODE_BG_COLOR = "F2F2F2"  # per style guide
# 提示块类型定义（来自中文 PRD 样式规范）
# 格式：左侧窄色条 + 右侧浅色背景内容区
CALLOUT_TYPES = {
    # marker: (stripe_color, bg_color, label)
    "[!INFO]":    ("2B5797", "E8F0FE", "信息提示"),
    "[!CAUTION]": ("E6A817", "FFF8E1", "注意事项"),
    "[!WARNING]": ("C0392B", "FDECEA", "风险警告"),
    "[!TIP]":     ("27AE60", "E8F5E9", "最佳实践"),
    # 自动映射 PRD 标记到提示块
    "[待确认]":   ("E6A817", "FFF8E1", "待确认"),
    "[推断]":     ("2B5797", "E8F0FE", "推断"),
    "[建议]":     ("27AE60", "E8F5E9", "建议"),
}
CALLOUT_STRIPE_WIDTH = 120  # DXA（色条宽度，与 JS 规范一致）
IMAGE_WIDTH = Inches(6.0)

# 表格样式常量（来自中文 PRD 样式规范）
TABLE_HEADER_BG = "2B5797"   # 蓝底
TABLE_HEADER_FG = "FFFFFF"   # 白字
TABLE_ZEBRA_BG = "F7F7F7"   # 偶数行底色
TABLE_BORDER_COLOR = "D0D0D0"  # 边框色
TABLE_FONT_SIZE = Pt(10)

# 排版常量（与 prd-docx-styles.docx 模板保持一致）
# 注意：标题/段落/列表间距由模板样式定义，不在代码中覆盖
CJK_FONT = "Microsoft YaHei"
CODE_SPACE = Pt(8)            # 代码块前后间距
ELEMENT_GAP = Pt(10)          # 元素间通用间隔（表格/提示块/图片前后，≈200 DXA）

# 表格排版（参考模板 §9.2 垂直间距节奏）
TABLE_CELL_H_MARGIN = 120     # 单元格水平内边距 (DXA)
TABLE_CELL_V_MARGIN = 80      # 单元格垂直内边距 (DXA)
TABLE_LINE_SPACING = 288      # 单元格行距 = 1.2 倍（240=单倍，288=1.2倍）
TABLE_SEQ_COL_WIDTH = 600     # 序号列宽度 (DXA, 约 1cm)

# H1 底部分割线
H1_BORDER_COLOR = "2B5797"    # 与 Heading 2 主题色一致

# 模板路径
_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), '..', 'templates', 'prd-docx-styles.docx')


# =============================================================================
# 内联格式处理
# =============================================================================

_INLINE_RE = re.compile(r"(\*\*.*?\*\*|`.*?`)")


def _set_run_shading(run, color_hex):
    """为 run 设置字符级背景色。"""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    run._element.get_or_add_rPr().append(shd)


def _set_cjk_font(run, font_name=None):
    """为 run 设置 CJK 字体。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name or CJK_FONT)


def _add_formatted_runs(paragraph, text):
    """将文本按 bold / code / normal 分段添加到段落中，统一设置 CJK 字体。"""
    segments = _INLINE_RE.split(text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
            _set_cjk_font(run)
        elif seg.startswith("`") and seg.endswith("`"):
            run = paragraph.add_run(seg[1:-1])
            run.font.name = MONOSPACE_FONT
            run.font.size = MONOSPACE_SIZE
            _set_cjk_font(run, MONOSPACE_FONT)
            _set_run_shading(run, CODE_BG_COLOR)
        else:
            run = paragraph.add_run(seg)
            _set_cjk_font(run)


def _add_formatted_runs_to_cell(cell, text):
    """为表格单元格添加格式化文本（清除默认空段落后写入）。"""
    # 单元格创建时自带一个空段落，复用它
    paragraph = cell.paragraphs[0]
    _add_formatted_runs(paragraph, text.strip())


# =============================================================================
# 图片辅助
# =============================================================================

def _resolve_image_path(raw_path, md_dir):
    """解析图片路径，尝试多种后缀变体。返回 (resolved_path | None, filename)。"""
    # 清理路径
    raw_path = raw_path.strip()
    filename = os.path.basename(raw_path)

    # 构造绝对路径
    if os.path.isabs(raw_path):
        candidate = raw_path
    else:
        candidate = os.path.normpath(os.path.join(md_dir, raw_path))

    if os.path.isfile(candidate):
        return candidate, filename

    # 如果是 .drawio / .svg / .mermaid，尝试同名 .png
    base, ext = os.path.splitext(candidate)
    if ext.lower() in (".drawio", ".svg", ".mermaid"):
        png_candidate = base + ".png"
        if os.path.isfile(png_candidate):
            return png_candidate, os.path.basename(png_candidate)

    # 在 diagrams/ 子目录中查找同名 png
    diagrams_dir = os.path.join(md_dir, "diagrams")
    if os.path.isdir(diagrams_dir):
        png_name = os.path.splitext(filename)[0] + ".png"
        diagrams_candidate = os.path.join(diagrams_dir, png_name)
        if os.path.isfile(diagrams_candidate):
            return diagrams_candidate, png_name

    return None, filename


# =============================================================================
# 代码块灰底辅助
# =============================================================================

def _set_paragraph_shading(paragraph, color_hex):
    """为段落设置背景底色。"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    paragraph._element.get_or_add_pPr().append(shading_elm)


# =============================================================================
# 有序列表编号重置
# =============================================================================

# 缓存 ListNumber 样式的 abstractNumId（文档级，只查一次）
_list_number_abstract_id = None
_next_num_id_counter = None


def _restart_list_numbering(paragraph):
    """强制当前段落的有序列表从 1 重新编号。

    通过在 Word numbering.xml 中创建新的 <w:num> 引用同一 abstractNumId
    但附加 <w:startOverride val="1"/>，然后在段落 pPr 中注入该 numId。
    """
    global _list_number_abstract_id, _next_num_id_counter

    numbering_el = paragraph.part.numbering_part.element

    # 首次调用：查找 ListNumber 样式的 abstractNumId 和最大 numId
    if _list_number_abstract_id is None:
        for abstract in numbering_el.findall(qn('w:abstractNum')):
            abs_id = abstract.get(qn('w:abstractNumId'))
            for lvl in abstract.findall(qn('w:lvl')):
                pStyle = lvl.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'ListNumber':
                    _list_number_abstract_id = abs_id
                    break
            if _list_number_abstract_id:
                break
        _next_num_id_counter = max(
            int(n.get(qn('w:numId')))
            for n in numbering_el.findall(qn('w:num'))
        ) + 1

    if _list_number_abstract_id is None:
        return  # 未找到 ListNumber 定义，跳过

    # 创建新的 num 元素
    new_id = _next_num_id_counter
    _next_num_id_counter += 1
    new_num = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="{new_id}">'
        f'  <w:abstractNumId w:val="{_list_number_abstract_id}"/>'
        f'  <w:lvlOverride w:ilvl="0">'
        f'    <w:startOverride w:val="1"/>'
        f'  </w:lvlOverride>'
        f'</w:num>'
    )
    numbering_el.append(new_num)

    # 在段落 pPr 中注入 numPr 覆盖样式继承的编号
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        paragraph._element.insert(0, pPr)
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="0"/>'
        f'  <w:numId w:val="{new_id}"/>'
        f'</w:numPr>'
    )
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is not None:
        pStyle.addnext(numPr)
    else:
        pPr.insert(0, numPr)


# =============================================================================
# 水平分割线
# =============================================================================

def _add_horizontal_rule(doc):
    """添加一条细水平线（使用底部边框模拟）。"""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    p.space_before = Pt(6)
    pPr = p._element.get_or_add_pPr()
    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="AAAAAA"/>'
        "</w:pBdr>"
    )
    pPr.append(parse_xml(border_xml))


# =============================================================================
# 主转换逻辑
# =============================================================================

# 正则模式
_RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_RE_IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|\s*$")
_RE_TABLE_SEP = re.compile(r"^\|[\s:]*-{2,}[\s:|-]*\|\s*$")
_RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)")
_RE_UL = re.compile(r"^(\s*)[-*]\s+(.*)")
_RE_OL = re.compile(r"^(\s*)\d+\.\s+(.*)")
_RE_CODE_FENCE = re.compile(r"^```(\w*)")
_RE_HR = re.compile(r"^---+\s*$")
_RE_FRONT_MATTER_DELIM = re.compile(r"^---\s*$")


def _apply_fallback_styles(doc):
    """模板不可用时的降级样式。"""
    try:
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    except KeyError:
        pass


def _set_cell_shading(cell, color_hex):
    """为表格单元格设置背景色。"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading_elm)


def _is_seq_column(col_values):
    """判断某列是否为序号列（内容为纯数字或 #、序号 等）。"""
    seq_patterns = re.compile(r'^(\d{1,3}|#|序号|No\.?)$')
    return all(seq_patterns.match(v.strip()) for v in col_values if v.strip())


def _compute_col_widths(table_rows, total_width_pct=5000):
    """根据内容计算列宽分配（百分比单位，5000 = 100%）。

    策略：
    - 序号列固定窄宽
    - 标题行文字权重最高（避免标题换行）
    - 其余列按最大内容长度加权分配
    """
    if not table_rows:
        return []
    num_cols = len(table_rows[0])
    if num_cols == 0:
        return []

    def char_width(text):
        """计算显示宽度：中文/全角算 2，ASCII 算 1。"""
        return sum(2 if ord(ch) > 0x7F else 1 for ch in text)

    # 判断序号列（第一列，数据行均为短数字/符号）
    is_seq = [False] * num_cols
    if num_cols > 1 and len(table_rows) > 1:
        data_vals = [row[0] for row in table_rows[1:]]
        if data_vals and _is_seq_column(data_vals):
            is_seq[0] = True

    # 计算每列所需宽度
    # 标题行权重 ×2（优先保证标题不换行）
    # 数据行取 P90 宽度（忽略极端长行），避免个别长文本拉宽整列
    col_needs = []
    for c_idx in range(num_cols):
        if is_seq[c_idx]:
            col_needs.append(0)
            continue
        header_w = char_width(table_rows[0][c_idx]) * 2.0 if table_rows[0] else 0
        data_widths = sorted(
            (char_width(row[c_idx]) for row in table_rows[1:] if c_idx < len(row)),
            reverse=True
        )
        # P90: 取第 10% 位置的值（忽略最宽的几行）
        if data_widths:
            p90_idx = max(0, len(data_widths) // 10)
            data_w = data_widths[min(p90_idx, len(data_widths) - 1)]
        else:
            data_w = 0
        col_needs.append(max(header_w, data_w, 3))

    # 序号列固定 ~6%
    seq_pct = 300
    seq_total = sum(seq_pct for s in is_seq if s)
    remaining = total_width_pct - seq_total

    total_need = sum(col_needs) or 1
    result = []
    for c_idx in range(num_cols):
        if is_seq[c_idx]:
            result.append(seq_pct)
        else:
            pct = int(remaining * col_needs[c_idx] / total_need)
            pct = max(350, min(3000, pct))  # 最小 7%, 最大 60%
            result.append(pct)

    # 归一化
    total = sum(result)
    if total != total_width_pct:
        factor = total_width_pct / total
        result = [int(w * factor) for w in result]
        diff = total_width_pct - sum(result)
        for c_idx in range(num_cols - 1, -1, -1):
            if not is_seq[c_idx]:
                result[c_idx] += diff
                break

    return result


def _style_table(tbl, table_rows=None):
    """应用表格样式：全宽 + 蓝底白字表头 + 斑马纹 + 边框 + 内边距 + 行距 + 列宽。"""
    tbl_pr = tbl._element.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl._element.insert(0, tbl_pr)

    # 表格 100% 页宽
    tbl_pr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))

    # 单元格默认内边距：上下 80 DXA、左右 120 DXA
    tbl_pr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{TABLE_CELL_V_MARGIN}" w:type="dxa"/>'
        f'  <w:left w:w="{TABLE_CELL_H_MARGIN}" w:type="dxa"/>'
        f'  <w:bottom w:w="{TABLE_CELL_V_MARGIN}" w:type="dxa"/>'
        f'  <w:right w:w="{TABLE_CELL_H_MARGIN}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))

    # 边框
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(borders_xml))

    # 列宽分配
    col_widths = _compute_col_widths(table_rows) if table_rows else []

    for r_idx, row in enumerate(tbl.rows):
        # 表头行：换页时重复
        if r_idx == 0:
            tr_pr = row._element.get_or_add_trPr()
            tr_pr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for c_idx, cell in enumerate(row.cells):
            tc_pr = cell._element.get_or_add_tcPr()

            # 垂直居中
            tc_pr.append(parse_xml(
                f'<w:vAlign {nsdecls("w")} w:val="center"/>'
            ))

            # 列宽
            if col_widths and c_idx < len(col_widths):
                tc_pr.append(parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{col_widths[c_idx]}" w:type="pct"/>'
                ))

            # 单元格内段落：1.2 倍行距 + 紧凑间距 + 字体
            for paragraph in cell.paragraphs:
                pPr = paragraph._element.get_or_add_pPr()
                pPr.append(parse_xml(
                    f'<w:spacing {nsdecls("w")} w:line="{TABLE_LINE_SPACING}"'
                    f' w:lineRule="auto" w:before="0" w:after="0"/>'
                ))
                for run in paragraph.runs:
                    run.font.size = TABLE_FONT_SIZE
                    _set_cjk_font(run)

            if r_idx == 0:
                # 表头行：蓝底白字加粗
                _set_cell_shading(cell, TABLE_HEADER_BG)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # 数据行：斑马纹
                if r_idx % 2 == 0:
                    _set_cell_shading(cell, TABLE_ZEBRA_BG)


def _detect_callout(text):
    """检测文本是否匹配提示块类型，返回 (marker, content) 或 None。"""
    for marker in CALLOUT_TYPES:
        if marker in text:
            content = text.replace(marker, "", 1).strip()
            return marker, content
    return None


def _add_callout_block(doc, marker, content):
    """渲染提示块：左侧窄色条 + 右侧浅色背景内容区（双列无边框表格）。"""
    stripe_color, bg_color, label = CALLOUT_TYPES[marker]

    # 提示块前间距：插入间隔段落（前一个元素可能是表格，无法用 space_after）
    spacer = doc.add_paragraph()
    spacer_run = spacer.add_run()
    spacer_run.font.size = Pt(1)
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = ELEMENT_GAP

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False

    tbl_pr = tbl._element.tblPr
    # 全宽
    tbl_pr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))
    # 隐藏所有边框
    no_borders = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(no_borders))

    # 左列：窄色条（仅用 tcW 控制宽度，零内边距）
    stripe_cell = tbl.cell(0, 0)
    tc_pr = stripe_cell._element.get_or_add_tcPr()
    tc_pr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{stripe_color}" w:val="clear"/>'
    ))
    tc_pr.append(parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{CALLOUT_STRIPE_WIDTH}" w:type="dxa"/>'
    ))
    tc_pr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="0" w:type="dxa"/><w:left w:w="0" w:type="dxa"/>'
        f'  <w:bottom w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/>'
        f'</w:tcMar>'
    ))
    stripe_cell.paragraphs[0].text = ""

    # 右列：浅色背景 + 内容 + 内边距
    content_cell = tbl.cell(0, 1)
    tc_pr2 = content_cell._element.get_or_add_tcPr()
    tc_pr2.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>'
    ))
    tc_pr2.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="100" w:type="dxa"/>'
        f'  <w:left w:w="180" w:type="dxa"/>'
        f'  <w:bottom w:w="100" w:type="dxa"/>'
        f'  <w:right w:w="160" w:type="dxa"/>'
        f'</w:tcMar>'
    ))
    # 内容：标签加粗 + 正文
    paragraphs = content.split("\n")
    first_para = True
    for para_text in paragraphs:
        if first_para:
            p = content_cell.paragraphs[0]
            label_run = p.add_run(f"{label}  ")
            label_run.bold = True
            label_run.font.size = TABLE_FONT_SIZE
            label_run.font.color.rgb = RGBColor.from_string(stripe_color)
            _set_cjk_font(label_run)
            first_para = False
        else:
            if not para_text.strip():
                continue
            p = content_cell.add_paragraph()
        _add_formatted_runs(p, para_text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)


def convert(md_path):
    """将 Markdown 文件转换为 .docx，返回输出文件路径。"""
    md_path = os.path.abspath(md_path)
    md_dir = os.path.dirname(md_path)

    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    lines = [l.rstrip("\n") for l in raw_lines]

    # 加载样式模板（继承字体/标题/页面设置），降级为空文档
    template = os.path.abspath(_TEMPLATE_PATH)
    if os.path.isfile(template):
        doc = Document(template)
        # 清空模板示例内容，只保留样式和页面设置（sectPr）
        body = doc.element.body
        for child in list(body):
            # 保留 sectPr（页面设置）
            if child.tag.endswith('}sectPr'):
                continue
            body.remove(child)
    else:
        doc = Document()
        _apply_fallback_styles(doc)

    # -----------------------------------------------------------------
    # 状态机
    # -----------------------------------------------------------------
    i = 0
    total = len(lines)
    h1_count = 0      # 跟踪 Word H1 数量，用于分页控制
    doc_title = ""    # 文档标题，用于页眉

    # 跳过 YAML front matter
    if total > 0 and _RE_FRONT_MATTER_DELIM.match(lines[0]):
        i = 1
        while i < total and not _RE_FRONT_MATTER_DELIM.match(lines[i]):
            i += 1
        i += 1  # 跳过结束 ---

    while i < total:
        line = lines[i]

        # ---- 空行 ----
        if line.strip() == "":
            i += 1
            continue

        # ---- 水平分割线 ----
        # PRD 中 --- 仅作为 Markdown 章节分隔，Word 中由 H1 分页替代，跳过
        if _RE_HR.match(line):
            i += 1
            continue

        # ---- 标题 ----
        # Markdown → Word 标题层级映射：
        #   # (md H1)  → Title 样式（文档总标题，26pt）
        #   ## (md H2) → Heading 1（章标题，18pt，分页+底线）
        #   ### (md H3) → Heading 2（节标题，14pt）
        #   #### (md H4) → Heading 3（小节标题，12pt）
        m_heading = _RE_HEADING.match(line)
        if m_heading:
            md_level = len(m_heading.group(1))
            text = m_heading.group(2).strip()

            if md_level == 1:
                # Markdown # → 封面页标题
                doc_title = text
                # 封面居中留白
                for _ in range(6):
                    sp = doc.add_paragraph()
                    sp.paragraph_format.space_after = Pt(0)
                # 标题：30pt 加粗，左侧蓝色竖线
                heading = doc.add_paragraph()
                heading.paragraph_format.space_after = Pt(4)
                pPr = heading._element.get_or_add_pPr()
                pPr.append(parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="20" w:space="8"'
                    f'   w:color="{H1_BORDER_COLOR}"/>'
                    f'</w:pBdr>'
                ))
                pPr.append(parse_xml(
                    f'<w:ind {nsdecls("w")} w:left="400"/>'
                ))
                title_run = heading.add_run(text)
                title_run.bold = True
                title_run.font.size = Pt(30)
                title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                _set_cjk_font(title_run)
                # 封面后分页
                pb = doc.add_paragraph()
                pb_run = pb.add_run()
                pb_run.add_break()  # page break
            else:
                # md ## → Word H1, md ### → Word H2, md #### → Word H3
                word_level = md_level - 1
                heading = doc.add_heading(level=word_level)
                _add_formatted_runs(heading, text)
                # 标题间距由模板样式定义，不在此覆盖
                if word_level == 1:
                    # 章标题（Word H1）：始终分页 + 底部分割线
                    heading.paragraph_format.page_break_before = True
                    h1_count += 1
                    pPr = heading._element.get_or_add_pPr()
                    pPr.append(parse_xml(
                        f'<w:pBdr {nsdecls("w")}>'
                        f'  <w:bottom w:val="single" w:sz="8" w:space="4"'
                        f'   w:color="{H1_BORDER_COLOR}"/>'
                        f'</w:pBdr>'
                    ))
            i += 1
            continue

        # ---- 图片 ----
        m_img = _RE_IMAGE.match(line)
        if m_img:
            alt_text = m_img.group(1)
            img_raw = m_img.group(2)
            resolved, filename = _resolve_image_path(img_raw, md_dir)
            if resolved:
                try:
                    doc.add_picture(resolved, width=IMAGE_WIDTH)
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p.paragraph_format.space_before = ELEMENT_GAP
                except Exception as exc:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[图表加载失败: {filename} — {exc}]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[图表: {filename}，见 diagrams/ 目录]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            if alt_text:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt_text)
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                cap.paragraph_format.space_after = ELEMENT_GAP
            i += 1
            continue

        # ---- 代码块 ----
        m_code = _RE_CODE_FENCE.match(line)
        if m_code:
            lang = m_code.group(1).lower()
            i += 1
            code_lines = []
            while i < total and not _RE_CODE_FENCE.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            if i < total:
                i += 1  # 跳过结束 ```

            # 图表 DSL（mermaid/plantuml/dot）：不逐行渲染源码，输出一行占位说明
            if lang in ("mermaid", "plantuml", "dot", "graphviz"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[{lang} 图表源码，见 diagrams/ 目录]")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            else:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                _set_paragraph_shading(p, CODE_BG_COLOR)
                p.space_before = CODE_SPACE
                p.space_after = CODE_SPACE
                run = p.add_run(code_text)
                run.font.name = MONOSPACE_FONT
                run.font.size = MONOSPACE_SIZE
                _set_cjk_font(run, MONOSPACE_FONT)
            continue

        # ---- 表格 ----
        if _RE_TABLE_ROW.match(line):
            table_rows = []
            while i < total and _RE_TABLE_ROW.match(lines[i]):
                if _RE_TABLE_SEP.match(lines[i]):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(cells)
                i += 1

            if not table_rows:
                continue

            # 统一列数——取最大列数
            max_cols = max(len(r) for r in table_rows)
            for row in table_rows:
                while len(row) < max_cols:
                    row.append("")

            num_rows = len(table_rows)
            tbl = doc.add_table(rows=num_rows, cols=max_cols)

            for r_idx, row_data in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tbl.cell(r_idx, c_idx)
                    _add_formatted_runs_to_cell(cell, cell_text)

            _style_table(tbl, table_rows)

            # 表格后间距：插入间隔段落
            spacer = doc.add_paragraph()
            spacer_run = spacer.add_run()
            spacer_run.font.size = Pt(1)
            spacer.paragraph_format.space_before = ELEMENT_GAP
            spacer.paragraph_format.space_after = Pt(0)

            continue

        # ---- 引用块 ----
        m_bq = _RE_BLOCKQUOTE.match(line)
        if m_bq:
            # 收集连续引用块，在空行处分割为独立块
            all_blocks = []
            current_block = []
            while i < total:
                m = _RE_BLOCKQUOTE.match(lines[i])
                if m:
                    current_block.append(m.group(1))
                    i += 1
                elif lines[i].strip() == "":
                    if i + 1 < total and _RE_BLOCKQUOTE.match(lines[i + 1]):
                        # 空行分隔：保存当前块，开始新块
                        if current_block:
                            all_blocks.append(current_block)
                            current_block = []
                        i += 1
                    else:
                        break
                else:
                    break
            if current_block:
                all_blocks.append(current_block)

            # 逐块渲染
            for block_lines in all_blocks:
                bq_text = "\n".join(block_lines)
                callout = _detect_callout(bq_text)
                if callout:
                    marker, content = callout
                    _add_callout_block(doc, marker, content)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    _add_formatted_runs(p, bq_text)
                    # 间距由 Normal 样式定义

            continue

        # ---- 无序列表 ----
        m_ul = _RE_UL.match(line)
        if m_ul:
            while i < total:
                m = _RE_UL.match(lines[i])
                if m:
                    text = m.group(2)
                    p = doc.add_paragraph(style="List Bullet")
                    _add_formatted_runs(p, text)
                    # 间距由模板 contextualSpacing 控制
                    i += 1
                elif lines[i].strip() == "":
                    if i + 1 < total and _RE_UL.match(lines[i + 1]):
                        i += 1
                    else:
                        break
                else:
                    break
            continue

        # ---- 有序列表 ----
        m_ol = _RE_OL.match(line)
        if m_ol:
            is_first_item = True
            while i < total:
                m = _RE_OL.match(lines[i])
                if m:
                    text = m.group(2)
                    p = doc.add_paragraph(style="List Number")
                    _add_formatted_runs(p, text)
                    if is_first_item:
                        _restart_list_numbering(p)
                        is_first_item = False
                    # 间距由模板 contextualSpacing 控制
                    i += 1
                elif lines[i].strip() == "":
                    if i + 1 < total and _RE_OL.match(lines[i + 1]):
                        i += 1
                    else:
                        break
                else:
                    break
            continue

        # ---- 普通段落 ----
        para_lines = [line]
        i += 1
        # 合并连续非空行（非特殊语法行）为一个段落
        while i < total:
            next_line = lines[i]
            if (
                next_line.strip() == ""
                or _RE_HEADING.match(next_line)
                or _RE_IMAGE.match(next_line)
                or _RE_CODE_FENCE.match(next_line)
                or _RE_TABLE_ROW.match(next_line)
                or _RE_BLOCKQUOTE.match(next_line)
                or _RE_UL.match(next_line)
                or _RE_OL.match(next_line)
                or _RE_HR.match(next_line)
            ):
                break
            para_lines.append(next_line)
            i += 1

        para_text = " ".join(para_lines)
        # 处理段落内的图片链接（行内）
        inline_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", para_text.strip())
        if inline_img:
            alt_text = inline_img.group(1)
            img_raw = inline_img.group(2)
            resolved, filename = _resolve_image_path(img_raw, md_dir)
            if resolved:
                try:
                    doc.add_picture(resolved, width=IMAGE_WIDTH)
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p.paragraph_format.space_before = ELEMENT_GAP
                except Exception:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[图表: {filename}，见 diagrams/ 目录]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[图表: {filename}，见 diagrams/ 目录]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, para_text)
            # 间距由 Normal 样式定义（space_after=6pt, line=1.5x）

    # -----------------------------------------------------------------
    # 页眉页脚
    # -----------------------------------------------------------------
    section = doc.sections[0]

    # 页眉：文档标题（右对齐，9pt 灰色）
    header = section.header
    header.is_linked_to_previous = False
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_para.add_run(doc_title or os.path.splitext(os.path.basename(md_path))[0])
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _set_cjk_font(h_run)

    # 页脚：页码居中（"第 X 页 / 共 Y 页"）
    footer = section.footer
    footer.is_linked_to_previous = False
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run1 = f_para.add_run("第 ")
    f_run1.font.size = Pt(9)
    f_run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # 当前页码域
    fld_page = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "/>'
    )
    r_page = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/>'
                       f'<w:color w:val="999999"/></w:rPr><w:t>1</w:t></w:r>')
    fld_page.append(r_page)
    f_para._element.append(fld_page)
    f_run2 = f_para.add_run(" 页 / 共 ")
    f_run2.font.size = Pt(9)
    f_run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # 总页数域
    fld_numpages = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "/>'
    )
    r_nump = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/>'
                       f'<w:color w:val="999999"/></w:rPr><w:t>1</w:t></w:r>')
    fld_numpages.append(r_nump)
    f_para._element.append(fld_numpages)
    f_run3 = f_para.add_run(" 页")
    f_run3.font.size = Pt(9)
    f_run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # -----------------------------------------------------------------
    # 保存
    # -----------------------------------------------------------------
    base, _ = os.path.splitext(md_path)
    output_path = base + ".docx"
    doc.save(output_path)
    return output_path


# =============================================================================
# CLI 入口
# =============================================================================

def main():
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <PRD.md文件路径>", file=sys.stderr)
        print("示例: python md2docx.py requirements/REQ-xxx/PRD-xxx.md", file=sys.stderr)
        sys.exit(1)

    md_path = sys.argv[1]

    if not os.path.isfile(md_path):
        print(f"错误: 文件不存在 — {md_path}", file=sys.stderr)
        sys.exit(1)

    output = convert(md_path)
    print(f"✓ 已生成: {output}")


if __name__ == "__main__":
    main()
